import json
import os
from typing import Dict, Any, Tuple

# Global flags to track module availability
PYPDF2_AVAILABLE = False
DOCX_AVAILABLE = False

# For PDF parsing
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    print("PyPDF2 not installed. PDF parsing will not be available.")

# For DOCX parsing
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    print("python-docx not installed. DOCX parsing will not be available.")

def parse_pdf(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """
    Parse a PDF file and extract its text content using PyMuPDF.
    Implements multiple extraction methods for better reliability.
    For scanned PDFs, uses OCR to extract text from page images.
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        Tuple[str, Dict[str, Any]]: Extracted text and metadata
    """
    try:
      import fitz  # PyMuPDF
    except ImportError:
        raise Exception("PyMuPDF is not installed. Please install it with 'pip install PyMuPDF'.")

    try:
        # Check for pytesseract for OCR
        ocr_available = False
        try:
            import pytesseract
            from PIL import Image
            import io
            ocr_available = True
        except ImportError:
            pass
            
        # Open the document
        doc = fitz.open(file_path)
        
        # Check if document is encrypted/password-protected
        if doc.is_encrypted:
            try:
                # Try opening with empty password
                success = doc.authenticate("")
                if not success:
                    raise Exception("PDF is password protected. Cannot proceed with extraction.")
            except Exception as e:
                raise Exception(f"PDF is encrypted: {str(e)}")
        
        # Save document properties we'll need later
        page_count = doc.page_count
        doc_metadata = doc.metadata.copy() if doc.metadata else {}
        doc_format = doc_metadata.get("format", "pdf")
        
        # Initialize text content
        text = ""
        problematic_pages = []
        is_scanned_pdf = True  # Assume it's scanned until we find actual text
        
        # First pass: Try standard text extraction methods
        for page_num in range(page_count):
            try:
                page = doc[page_num]
                
                # Try different text extraction methods in order of reliability
                page_text = page.get_text().strip()
                
                # If we found substantial text, it's not a scanned PDF
                if len(page_text) > 100:
                    is_scanned_pdf = False
                    
                # If default extraction returns empty or very short text, try other methods
                if not page_text or len(page_text) < 50:
                    # Try extracting with the "text" option for better handling of tables
                    text_option = page.get_text("text").strip()
                    if len(text_option) > len(page_text):
                        page_text = text_option
                    
                    # If still not getting good results, try HTML mode
                    if not page_text or len(page_text) < 50:
                        html_text = page.get_text("html")
                        import re
                        html_option = re.sub(r'<[^>]*>', ' ', html_text).strip()
                        if len(html_option) > len(page_text):
                            page_text = html_option
                
                    # If still no text, try extracting as blocks
                    if not page_text or len(page_text) < 50:
                        blocks = page.get_text("blocks")
                        blocks_text = "\n".join([b[4] for b in blocks if len(b) > 4]).strip()
                        if len(blocks_text) > len(page_text):
                            page_text = blocks_text
                            
                text += page_text + "\n"
                
            except Exception as e:
                problematic_pages.append(page_num)
                print(f"Warning: Could not extract text from page {page_num+1}: {str(e)}")
                continue
        
        # Check if we need OCR
        need_ocr = (len(text.strip()) < 100 or is_scanned_pdf) and page_count > 0
        
        # Second pass: If little text was found, it's likely a scanned PDF, so use OCR
        if need_ocr:
            print(f"Detected likely scanned PDF with minimal text. Attempting OCR...")
            
            if ocr_available:
                ocr_text = ""  # Separate variable for OCR results
                
                for page_num in range(page_count):
                    try:
                        page = doc[page_num]
                        
                        # Higher DPI for better OCR results
                        pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
                        
                        # Convert pixmap to PIL Image for OCR
                        img_bytes = pix.tobytes("png")
                        img = Image.open(io.BytesIO(img_bytes))
                        
                        # Perform OCR
                        page_text = pytesseract.image_to_string(img)
                        ocr_text += page_text + "\n"
                        
                    except Exception as e:
                        problematic_pages.append(page_num)
                        print(f"Warning: OCR failed on page {page_num+1}: {str(e)}")
                
                # If OCR produced content, use it
                if ocr_text.strip():
                    text = ocr_text
            else:
                # No OCR available, provide a helpful message
                print("Warning: This appears to be a scanned PDF. Install pytesseract for OCR capabilities.")
                # Try PyMuPDF's built-in alternative for image-based text detection
                try:
                    xhtml_text = ""
                    for page_num in range(page_count):
                        try:
                            page = doc[page_num]
                            # Attempt extracting image text using XHTML which might include alt text
                            xhtml = page.get_text("xhtml")
                            import re
                            clean_text = re.sub(r'<[^>]*>', ' ', xhtml).strip()
                            xhtml_text += clean_text + "\n"
                        except Exception as e:
                            continue
                    
                    # If XHTML produced more content, use it
                    if len(xhtml_text.strip()) > len(text.strip()):
                        text = xhtml_text
                except Exception as e:
                    print(f"Warning: Alternative text extraction failed: {str(e)}")
        
        # Create metadata dictionary
        metadata = {
            "filename": os.path.basename(file_path),
            "pages": page_count,
            "format": doc_format,
            "problematic_pages": problematic_pages,
            "is_scanned_pdf": is_scanned_pdf
        }
        
        # Add metadata from document
        if doc_metadata:
            for key, value in doc_metadata.items():
                if value and str(value).strip():
                    metadata[key.lower()] = str(value)
        
        # Now it's safe to close the document since we've extracted all we need
        doc.close()
        
        # Final check for empty content (using stored page_count)
        final_text = text.strip()
        if not final_text and page_count > 0:
            if is_scanned_pdf:
                if ocr_available:
                    return "Document appears to be scanned but OCR extraction failed. Try with a different PDF.", metadata
                else:
                    return "Document appears to be scanned. Install Tesseract OCR for better text extraction.", metadata
            else:
                return "Unable to extract text from this PDF. The file may be damaged or contain only images.", metadata
            
        return final_text, metadata
        
    except Exception as e:
        # Make sure document is closed if something went wrong
        if 'doc' in locals() and not isinstance(doc, Exception):
            try:
                doc.close()
            except:
                pass
                
        # Provide detailed error message
        import traceback
        error_details = traceback.format_exc()
        raise Exception(f"Error parsing PDF file: {str(e)}\nDetails: {error_details}")


def parse_docx(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """
    Parse a DOCX file and extract its text content.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        Tuple[str, Dict[str, Any]]: Extracted text and metadata
    """
    if not DOCX_AVAILABLE:
        raise Exception("python-docx library is not installed. Please install it with 'pip install python-docx'.")
        
    try:
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Extract basic metadata
        metadata = {
            "filename": os.path.basename(file_path),
            "paragraphs": len(doc.paragraphs),
            "format": "docx"
        }
        
        # Try to extract core properties if available
        if hasattr(doc, 'core_properties'):
            props = doc.core_properties
            metadata_fields = [
                'author', 'category', 'comments', 'content_status',
                'created', 'identifier', 'keywords', 'language',
                'last_modified_by', 'last_printed', 'modified',
                'revision', 'subject', 'title', 'version'
            ]
            
            for field in metadata_fields:
                if hasattr(props, field):
                    value = getattr(props, field)
                    if value:
                        metadata[field] = str(value)
        
        return text.strip(), metadata
    except Exception as e:
        raise Exception(f"Error parsing DOCX file: {str(e)}")


def parse_json(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """
    Parse a JSON file and convert it to text.
    
    Args:
        file_path (str): Path to the JSON file
        
    Returns:
        Tuple[str, Dict[str, Any]]: JSON as text and metadata
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
            
            # Convert JSON to formatted string
            text = json.dumps(data, indent=2)
            
            # Extract basic metadata
            metadata = {
                "filename": os.path.basename(file_path),
                "format": "json"
            }
            
            # If the JSON has metadata fields, extract them
            if isinstance(data, dict):
                if "metadata" in data:
                    metadata["document_metadata"] = str(data["metadata"])
                if "title" in data:
                    metadata["title"] = data["title"]
                if "author" in data:
                    metadata["author"] = data["author"]
            
            return text.strip(), metadata
    except Exception as e:
        raise Exception(f"Error parsing JSON file: {str(e)}")


def parse_txt(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """
    Parse a TXT file and extract its content.
    
    Args:
        file_path (str): Path to the TXT file
        
    Returns:
        Tuple[str, Dict[str, Any]]: Text content and metadata
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
            
            # Extract basic metadata
            metadata = {
                "filename": os.path.basename(file_path),
                "format": "txt",
                "size_bytes": os.path.getsize(file_path)
            }
            
            return text.strip(), metadata
    except Exception as e:
        raise Exception(f"Error parsing TXT file: {str(e)}")


def parse_document(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """
    Parse a document based on its file extension.
    
    Args:
        file_path (str): Path to the document
        
    Returns:
        Tuple[str, Dict[str, Any]]: Extracted text and metadata
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return parse_pdf(file_path)
    elif file_extension == '.docx':
        return parse_docx(file_path)
    elif file_extension == '.json':
        return parse_json(file_path)
    elif file_extension == '.txt':
        return parse_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}") 